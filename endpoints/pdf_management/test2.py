from docx import Document
from fpdf import FPDF

# Open the word document
document = Document("original.docx")

# Create a new pdf document
pdf = FPDF()
pdf.add_font('Arial Unicode', '', 'Arial Unicode MS Font.ttf')
pdf.set_font("Arial Unicode", size=12)
pdf.add_page()

# Extract the header and footer
header = document.sections[0].header
footer = document.sections[0].footer
# print(document.sections[0].document.paragraphs[0].text)
# Add the header and footer to the pdf
if header is not None:
    x = 0
    for para in header.paragraphs:
        text = para.text
        pdf.cell(x, 0+10, txt=text, align='L')
        x += 1
        print(f"{text}")
        # alignment = para.alignment
        # lines = text.splitlines()
        # for line in lines:
        #     print(line)
        #     if alignment == 'center':
        #         pdf.cell(0, 10, txt=line, align='C')
        #     elif alignment == 'right':
        #         pdf.cell(0, 10, txt=line, align='R')
        #     else:
        #         pdf.cell(0, 10, txt=line, align='L')
    # pdf.ln()

# Iterate through the paragraphs in the word document
for para in document.paragraphs:
    # Extract the text and formatting of the paragraph
    text = para.text
    text = text.replace("offer_to", "new text")
    alignment = para.alignment
    # Split the text into lines
    lines = text.splitlines()
    pdf.set_xy(0, pdf.get_y()+10)
    for line in lines:
        if alignment == 'center':
            pdf.cell(0, 10, txt=line, align='C')
        elif alignment == 'right':
            pdf.cell(0, 10, txt=line, align='R')
        else:
            pdf.cell(0, 10, txt=line, align='L')
    pdf.ln()


pdf.output("modifiedtest2.pdf")